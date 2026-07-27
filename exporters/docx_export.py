import io
from docx import Document
from docx.shared import RGBColor, Pt


def export_docx(
    grid: list[str], get_color_rgb, font_name: str = "Consolas", font_size: int = 9
) -> io.BytesIO:
    doc = Document()

    for section in doc.sections:
        section.left_margin = Pt(24)
        section.right_margin = Pt(24)
        section.top_margin = Pt(24)
        section.bottom_margin = Pt(24)

    for row_idx, row in enumerate(grid):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0

        if row.strip() == "":
            paragraph.add_run("\u00a0")
            continue

        for col_idx, ch in enumerate(row):
            run = paragraph.add_run(ch if ch != " " else "\u00a0")
            run.font.name = font_name
            run.font.size = Pt(font_size)
            if ch != " ":
                r, g, b = get_color_rgb(row_idx, col_idx)
                run.font.color.rgb = RGBColor(r, g, b)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
